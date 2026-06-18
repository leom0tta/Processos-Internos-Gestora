#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gerador de Laudo de Carteira - Versão 2.0
Consome API Gorila e compara com perfil de suitability
"""

import json
import requests
import sys
import os
from typing import Dict, List, Optional, Tuple
from datetime import datetime
from pathlib import Path
import logging
from dotenv import load_dotenv

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Erro: python-docx não instalado.")
    print("Instale com: pip install -r requirements.txt")
    sys.exit(1)

# ============================================================================
# CONFIGURAÇÃO DE LOGGING
# ============================================================================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ============================================================================
# CLASSE PRINCIPAL
# ============================================================================

class GorilaLaudo:
    """Classe para gerar laudo de carteira usando dados da API Gorila"""

    def __init__(self, env_path: str = ".env", asset_type_mapping: dict = None, asset_name_mapping: dict = None):
        # Carrega .env apenas se existir (local). No Render, as variáveis já estão no ambiente.
        env_file = Path(env_path)
        if env_file.exists():
            load_dotenv(env_path)
        else:
            logger.info("Arquivo .env não encontrado — usando variáveis de ambiente do sistema.")

        # Validar credenciais
        self.api_key = os.getenv('GORILA_API_KEY')
        if not self.api_key or self.api_key == 'sua_api_key_aqui':
            logger.error("GORILA_API_KEY não configurada")
            raise RuntimeError("GORILA_API_KEY não configurada. Configure a variável de ambiente.")

        # Carregar configurações do .env
        self.base_url = os.getenv('GORILA_API_BASE_URL', 'https://core.gorila.com.br')
        self.timeout = int(os.getenv('GORILA_API_TIMEOUT', '30'))

        # Carregar informações da empresa
        self.company_name = os.getenv('COMPANY_NAME', 'Fatorial Capital')
        self.company_phone = os.getenv('COMPANY_PHONE', '+55 11 XXXX-XXXX')
        self.company_email = os.getenv('COMPANY_EMAIL', 'contato@fatorial.com.br')

        # Carregar JSONs de configuração (sempre relativos ao próprio script)
        _base = Path(__file__).parent
        self.mapping_path = _base / "asset_type_mapping.json"
        self.suitability_path = _base / "suitability_profiles.json"

        # Se mappings foram injetados (ex: do BD), usa eles; senão lê do JSON local
        if asset_type_mapping is not None:
            self.asset_type_mapping = asset_type_mapping
        else:
            self.asset_type_mapping = self._load_json(self.mapping_path)
        self.suitability_profiles = self._load_json(self.suitability_path)

        # Mapeamentos por nome de ativo (prioridade sobre security_type)
        self.asset_name_mapping = asset_name_mapping if asset_name_mapping is not None else {'mappings': {}}

        # Carregar guia de fundos para liquidação
        # Aceita variável de ambiente GUIA_FUNDOS_PATH; se não definida, usa pasta guia-de-fundos/ relativa ao script
        guia_env = os.getenv('GUIA_FUNDOS_PATH', '')
        if guia_env:
            self.guia_fundos_path = Path(guia_env)
        else:
            self.guia_fundos_path = Path(__file__).parent / "guia-de-fundos" / "Guia-de-Fundos-Maio-2026.xlsx"
        self.mapa_liquidacao = self._carregar_guia_fundos()

        logger.info("Sistema inicializado com sucesso")

    @staticmethod
    def _load_json(path: Path) -> dict:
        """Carrega arquivo JSON com tratamento de erro"""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            logger.error(f"Arquivo não encontrado: {path}")
            raise
        except json.JSONDecodeError as e:
            logger.error(f"Erro ao decodificar JSON {path}: {e}")
            raise

    def _carregar_guia_fundos(self) -> dict:
        """Carrega guia de fundos e cria mapa de liquidação por nome do fundo"""
        import pandas as pd

        mapa = {}

        if not self.guia_fundos_path.exists():
            logger.warning(f"Guia de fundos não encontrado: {self.guia_fundos_path}")
            return mapa

        try:
            # Abas a processar
            abas = [
                'Fundos XP Locais',
                'Fundos XP Offshore Digital',
                'Fundos Cetipados',
                'Fundos Alternativos',
                'Fundos PR+'
            ]

            for aba in abas:
                try:
                    # Ler pulando 3 linhas conforme indicado
                    df = pd.read_excel(
                        str(self.guia_fundos_path),
                        sheet_name=aba,
                        skiprows=3
                    )

                    # Extrair colunas relevantes
                    if 'CNPJ' in df.columns and 'Liquidação' in df.columns:
                        for _, row in df.iterrows():
                            fundo_cnpj = row.get('CNPJ')
                            liquidacao = row.get('Liquidação')

                            if pd.notna(fundo_cnpj) and pd.notna(liquidacao):
                                # Normaliza CNPJ: remove pontuação e zero-preenche para 14 dígitos
                                cnpj_normalizado = ''.join(filter(str.isdigit, str(fundo_cnpj))).zfill(14)
                                if cnpj_normalizado:
                                    mapa[cnpj_normalizado] = str(liquidacao).strip()

                        logger.info(f"Carregados {len(df)} fundos da aba '{aba}'")
                except Exception as e:
                    logger.warning(f"Erro ao processar aba '{aba}': {e}")

            logger.info(f"Total de fundos mapeados: {len(mapa)}")
            return mapa

        except Exception as e:
            logger.error(f"Erro ao carregar guia de fundos: {e}")
            return mapa

    def _fazer_request(self, endpoint: str, params: dict = None) -> dict:
        """Faz request à API Gorila com tratamento de erros"""
        headers = {
            'authorization': f'{self.api_key}',
            'Content-Type': 'application/json'
        }

        url = f"{self.base_url}{endpoint}"

        try:
            response = requests.get(
                url,
                headers=headers,
                params=params,
                timeout=self.timeout
            )
            response.raise_for_status()
            return response.json()
        except requests.exceptions.HTTPError as e:
            if response.status_code == 401:
                logger.error("Erro de autenticação. Verifique sua API Key.")
            elif response.status_code == 404:
                logger.error(f"Recurso não encontrado: {url}")
            else:
                logger.error(f"Erro HTTP {response.status_code}: {e}")
            raise
        except requests.exceptions.RequestException as e:
            logger.error(f"Erro na requisição: {e}")
            raise

    def buscar_portfolios(self) -> List[dict]:
        """Lista todos os portfolios do cliente"""
        logger.info("Buscando portfolios...")
        data = self._fazer_request("/portfolios")
        portfolios = data.get('records', [])
        logger.info(f"Encontrados {len(portfolios)} portfolios")
        return portfolios

    def buscar_posicoes(self, portfolio_id: str) -> List[dict]:
        """Busca todas as posições de um portfolio"""
        logger.info(f"Buscando posições do portfolio {portfolio_id}...")
        try:
            data = self._fazer_request(
                f"/portfolios/{portfolio_id}/positions",
                {"limit": 1000}
            )
            posicoes = data.get('records', [])
            logger.info(f"Encontradas {len(posicoes)} posições")
            return posicoes
        except Exception as e:
            logger.error(f"Erro ao buscar posições: {e}")
            return []

    def buscar_valores_mercado(self, portfolio_id: str) -> Dict[str, float]:
        """Busca valores de mercado das posições"""
        logger.info(f"Buscando valores de mercado...")
        try:
            data = self._fazer_request(
                f"/portfolios/{portfolio_id}/positions/market-values",
                {"limit": 1000}
            )
            resultados = data.get('records', [])

            valores = {}
            for item in resultados:
                security_id = item.get('security', {}).get('id')
                valor = item.get('marketValue', 0)
                if security_id:
                    valores[security_id] = valor

            return valores
        except Exception as e:
            logger.error(f"Erro ao buscar valores de mercado: {e}")
            return {}

    # Mapeamento assetClass → classe para ativos do tipo FUNDQUOTE.
    # A Gorila usa FUNDQUOTE como type genérico para fundos; o assetClass
    # é o campo informativo real nesses casos.
    FUNDQUOTE_ASSETCLASS_MAP = {
        'FIXED_INCOME': 'Fundos de renda fixa ativa / crédito',
        'MULTIMARKET':  'Fundos multimercado / macro',
        'STOCKS':       'Renda variável (fundos, ações, FIIs, ETFs, exceto de renda fixa)',
        'OFFSHORE':     'Ativos no exterior',
        'CRYPTO':       'Criptoativos',
        'CASH':         'Saldo em Conta',
        'CURRENCY':     'Ativos no exterior',
        'TANGIBLE':     'Renda variável (fundos, ações, FIIs, ETFs, exceto de renda fixa)',
    }

    def mapear_tipo_para_classe(self, security_type: str, security_name: str,
                                gorila_asset_class: str = '') -> Tuple[str, bool]:
        """
        Mapeia ativo para classe, com prioridade:
          1. Nome do ativo (asset_name_mapping)
          2. Se type == FUNDQUOTE → usa assetClass via FUNDQUOTE_ASSETCLASS_MAP
          3. security_type mapping (asset_type_mapping)
        Retorna (classe, foi_encontrado)
        """
        # Prioridade 1: nome exato do ativo
        name_mappings = self.asset_name_mapping.get('mappings', {})
        if security_name and security_name in name_mappings:
            return name_mappings[security_name], True

        # Prioridade 2: FUNDQUOTE → classifica pelo assetClass da Gorila
        if security_type == 'FUNDQUOTE' and gorila_asset_class:
            classe = self.FUNDQUOTE_ASSETCLASS_MAP.get(gorila_asset_class)
            if classe:
                logger.info(f"FUNDQUOTE classificado por assetClass '{gorila_asset_class}': {security_name} → {classe}")
                return classe, True

        # Prioridade 3: security_type mapping
        type_mappings = self.asset_type_mapping.get('mappings', {})
        if security_type in type_mappings:
            return type_mappings[security_type], True

        security_type_upper = security_type.upper()
        for tipo, classe in type_mappings.items():
            if tipo.upper() == security_type_upper:
                return classe, True

        logger.warning(f"Ativo não mapeado: {security_type} ({security_name})")
        return None, False

    # Tipos que representam fundos (sujeitos à regra de liquidez D+0 → DI)
    _FUND_TYPES = {
        'FUNDQUOTE', 'DI_FUND', 'FIXED_INCOME_FUND', 'CREDIT_FUND',
        'MULTIMARKET_FUND', 'MACRO_FUND', 'HEDGE_FUND', 'EQUITY_FUND',
        'PRIVATE_EQUITY_FUND', 'INFRASTRUCTURE_FUND', 'REAL_ESTATE_FUND',
        'INTERNATIONAL_FUND', 'CRYPTO_FUND', 'FIDC',
    }

    def _is_fund_type(self, security_type: str) -> bool:
        return security_type.upper() in self._FUND_TYPES or 'FUND' in security_type.upper()

    def _is_liquidez_d0(self, security: dict) -> bool:
        """Retorna True apenas se o ativo tiver liquidez D+0 no guia de fundos."""
        cnpj_raw = security.get('cnpj', '') or security.get('taxId', '') or ''
        cnpj = ''.join(filter(str.isdigit, str(cnpj_raw))).zfill(14)
        if not cnpj or cnpj not in self.mapa_liquidacao:
            return False
        return self.mapa_liquidacao[cnpj].strip().upper().startswith('D+0')

    def pedir_mapeamento_usuario(self, security_type: str, security_name: str, classes_disponiveis: List[str]) -> str:
        """Pede ao usuário para mapear um novo tipo de segurança"""
        print("\n" + "="*80)
        print(f"⚠️  TIPO DE SEGURANÇA NÃO MAPEADO")
        print("="*80)
        print(f"Type: {security_type}")
        print(f"Nome: {security_name}\n")

        print("Classes de ativo disponíveis:")
        for idx, classe in enumerate(classes_disponiveis, 1):
            print(f"  {idx}. {classe}")

        while True:
            try:
                escolha = input(f"\nEscolha o número (1-{len(classes_disponiveis)}) ou 'skip' para ignorar: ").strip()

                if escolha.lower() == 'skip':
                    return None

                idx = int(escolha) - 1
                if 0 <= idx < len(classes_disponiveis):
                    classe_escolhida = classes_disponiveis[idx]

                    self.asset_type_mapping['mappings'][security_type] = classe_escolhida
                    with open(self.mapping_path, 'w', encoding='utf-8') as f:
                        json.dump(self.asset_type_mapping, f, indent=2, ensure_ascii=False)

                    logger.info(f"Mapeamento salvo: {security_type} -> {classe_escolhida}")
                    return classe_escolhida
                else:
                    print("Opção inválida. Tente novamente.")
            except ValueError:
                print("Entrada inválida. Digite um número ou 'skip'.")

    def processar_posicoes(
        self,
        posicoes: List[dict],
        valores_mercado: dict,
        perfil: str,
        mapeamentos_extras: dict = None,
        mapeamentos_extras_nome: dict = None,
    ) -> Tuple[List[dict], list]:
        """
        Processa posições e realiza mapeamento com prioridade:
          1. asset_name_mapping (BD, por nome)
          2. mapeamentos_extras_nome (manual do usuário, por nome)
          3. asset_type_mapping (BD, por tipo)
          4. mapeamentos_extras (manual do usuário, por tipo)

        Retorna (posicoes_processadas, ativos_nao_mapeados)
            ativos_nao_mapeados: lista de dicts {security_type, security_name}
                para ativos sem mapeamento em nenhuma fonte.
            Se ativos_nao_mapeados for não-vazio, posicoes_processadas estará incompleta
            e o chamador deve solicitar os mapeamentos ao usuário e chamar novamente.
        """
        if mapeamentos_extras is None:
            mapeamentos_extras = {}
        if mapeamentos_extras_nome is None:
            mapeamentos_extras_nome = {}

        posicoes_processadas = []
        ativos_nao_mapeados  = []
        vistos               = set()   # evita duplicatas (chave: security_name ou security_type)

        for posicao in posicoes:
            security           = posicao.get('security', {})
            security_type      = security.get('type', 'UNKNOWN')
            security_name      = security.get('name', 'N/A')
            security_id        = security.get('id')
            gorila_asset_class = security.get('assetClass', '')

            # Flag: classificação veio de name mapping (não deve ser refinada por liquidez)
            from_name_mapping = bool(
                security_name and
                security_name in self.asset_name_mapping.get('mappings', {})
            )

            # Classificação principal: nome → FUNDQUOTE/assetClass → tipo
            classe, foi_encontrado = self.mapear_tipo_para_classe(
                security_type, security_name, gorila_asset_class
            )

            # Refinamento: fundo com liquidez D+0 → DI (exceto quando veio de name mapping)
            if foi_encontrado and not from_name_mapping and self._is_fund_type(security_type):
                if self._is_liquidez_d0(security):
                    logger.info(f"Fundo D+0 → DI: {security_name}")
                    classe = 'Fundos de renda fixa referenciados DI'

            if not foi_encontrado:
                # Fallback 1: mapeamento manual por nome (fornecido pelo usuário nesta chamada)
                if security_name in mapeamentos_extras_nome:
                    classe = mapeamentos_extras_nome[security_name]
                    self.asset_name_mapping['mappings'][security_name] = classe
                    logger.info(f"Mapeamento manual por nome aplicado: {security_name} -> {classe}")

                # Fallback 2: mapeamento manual por tipo
                elif security_type in mapeamentos_extras:
                    classe = mapeamentos_extras[security_type]
                    self.asset_type_mapping['mappings'][security_type] = classe
                    logger.info(f"Mapeamento manual por tipo aplicado: {security_type} -> {classe}")

                else:
                    chave = security_name if security_name != 'N/A' else security_type
                    if chave not in vistos:
                        ativos_nao_mapeados.append({
                            'security_type': security_type,
                            'security_name': security_name,
                        })
                        vistos.add(chave)
                    continue  # pula a posição até receber o mapeamento

            valor_mercado = valores_mercado.get(security_id, 0)

            posicao['classe_ativo']  = classe
            posicao['valor_mercado'] = valor_mercado

            posicoes_processadas.append(posicao)

        return posicoes_processadas, ativos_nao_mapeados

    def calcular_alocacoes(self, posicoes: List[dict]) -> Tuple[Dict[str, float], float]:
        """Calcula alocação por classe de ativo e patrimônio total"""
        alocacoes = {}
        patrimonio_total = 0

        for posicao in posicoes:
            valor = posicao.get('valor_mercado', 0)
            patrimonio_total += valor

        for posicao in posicoes:
            classe = posicao.get('classe_ativo', 'Desconhecido')
            valor = posicao.get('valor_mercado', 0)

            if patrimonio_total > 0:
                pct = valor / patrimonio_total
            else:
                pct = 0

            if classe not in alocacoes:
                alocacoes[classe] = 0

            alocacoes[classe] += pct

        return patrimonio_total, alocacoes

    def analisar_suitability(self, alocacoes: Dict[str, float], perfil: str) -> List[dict]:
        """Compara alocações com limites do perfil"""
        limites = self.suitability_profiles[perfil]
        analise = []

        for classe, limite in limites.items():
            alocacao_atual = alocacoes.get(classe, 0)
            dentro_perfil = limite['min'] <= alocacao_atual <= limite['max']

            analise.append({
                'classe': classe,
                'alocacao_atual': alocacao_atual,
                'limite_min': limite['min'],
                'limite_max': limite['max'],
                'dentro_perfil': dentro_perfil,
                'obs': limite['obs']
            })

        return analise

    def classificar_liquidez(self, security: dict) -> str:
        """
        Classifica a liquidez de um ativo baseado em:
        1. Guia de fundos (se for um fundo)
        2. Tipo e vencimento (fallback)
        Retorna categoria de liquidez.
        """
        security_name = security.get('name', '').strip()
        asset_class = security.get('assetClass', '')
        security_type = security.get('type', '')
        maturity_date = security.get('maturityDate')

        # Normaliza CNPJ: aceita tanto com pontuação quanto só dígitos; zero-preenche para 14 dígitos
        _cnpj_raw = security.get('cnpj', '') or security.get('taxId', '') or ''
        security_cnpj = ''.join(filter(str.isdigit, str(_cnpj_raw))).zfill(14)

        # PRIMEIRO: Verificar no guia de fundos
        if security_cnpj and security_cnpj in self.mapa_liquidacao:
            liquidacao_guia = self.mapa_liquidacao[security_cnpj]

            # Mapear liquidação do guia para categorias padrão
            liquidacao_upper = liquidacao_guia.upper()

            if liquidacao_upper.startswith('D+0'):
                return "Disponível (D+0 / D+1)"
            elif liquidacao_upper.startswith('D+1'):
                return "Disponível (D+0 / D+1)"
            elif liquidacao_upper.startswith('D+2') or liquidacao_upper.startswith('D+3'):
                return "Curto Prazo (D+2 a D+30)"
            elif liquidacao_upper.startswith('D+') and any(c.isdigit() for c in liquidacao_upper):
                # D+4, D+5, D+6 em diante
                return "Médio Prazo (D+31 a D+90)"
            elif liquidacao_upper.startswith('D+30'):
                return "Médio Prazo (D+31 a D+90)"
            elif liquidacao_upper == '-':
                return "Sem liquidez (ilíquido)"

        # FALLBACK: Lógica baseada em tipo de ativo

        # Tangíveis/Custom (prioridade sobre renda fixa)
        if asset_class in ['TANGIBLE'] or 'CUSTOM' in security_type.upper():
            return "Sem liquidez (ilíquido)"

        # Ações e ETFs B3 (RV): D+2
        if asset_class in ['STOCKS'] or 'STOCK' in security_type.upper() or 'ETF' in security_type.upper():
            return "Curto Prazo (D+2 a D+30)"

        # Renda fixa com vencimento
        if asset_class in ['FIXED_INCOME'] or 'FIXED' in security_type.upper() or 'DI' in security_type.upper():
            if maturity_date:
                try:
                    mat_date = datetime.strptime(maturity_date, '%Y-%m-%d')
                    dias = (mat_date - datetime.now()).days
                    if dias <= 1:
                        return "Disponível (D+0 / D+1)"
                    elif dias <= 30:
                        return "Curto Prazo (D+2 a D+30)"
                    elif dias <= 90:
                        return "Médio Prazo (D+31 a D+90)"
                    else:
                        return "RF — vencimento determinado"
                except (ValueError, TypeError):
                    pass
            return "RF — vencimento determinado"

        # Fundos
        if 'FUND' in security_type.upper() or 'PE' in security_type.upper() or 'FIP' in security_type.upper():
            if 'PE' in security_type.upper() or 'PRIVATE' in security_type.upper():
                return "Sem liquidez (ilíquido)"
            if 'DI' in security_type.upper() or 'FIXED_INCOME' in security_type.upper():
                return "Disponível (D+0 / D+1)"
            return "Médio Prazo (D+31 a D+90)"

        # FII
        if 'FII' in security_type.upper():
            return "Disponível (D+0 / D+1)"

        # FIDC
        if 'FIDC' in security_type.upper():
            return "Sem liquidez (ilíquido)"

        return "Médio Prazo (D+31 a D+90)"

    def gerar_docx(self, cliente_nome: str, perfil: str, alocacoes: Dict[str, float],
                    patrimonio_total: float, posicoes: List[dict], analise_suit: List[dict],
                    output_path: str = None) -> str:
        """Gera documento DOCX com o laudo"""

        if output_path is None:
            data_str = datetime.now().strftime("%d%m%Y")
            output_path = f"Laudo_{cliente_nome.replace(' ', '_')}_{data_str}.docx"

        doc = Document()
        self._set_document_font(doc, 'Arial')

        # ===== HEADER =====
        self._add_header(doc)

        # ===== TÍTULO =====
        title_table = doc.add_table(rows=1, cols=1)
        self._set_table_full_width(title_table)
        title_cell = title_table.cell(0, 0)
        self._set_cell_background(title_cell, '1B3A5C')

        # Padding generoso para a linha ficar maior que o texto
        self._set_table_cell_margins(title_table, top=90, left=160, bottom=90, right=160)

        title_para = title_cell.paragraphs[0]
        title_para.paragraph_format.line_spacing = 1
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(0)
        title_run = title_para.add_run("LAUDO DE ADERÊNCIA AO SUITABILITY")
        title_run.font.bold = True
        title_run.font.size = Pt(13)
        title_run.font.color.rgb = RGBColor(255, 255, 255)
        title_run.font.name = 'Arial'


        # ===== DADOS DO CLIENTE =====
        doc.add_paragraph()
        self._add_section_title(doc, "1. INFORMAÇÕES DO CLIENTE")

        info_table = doc.add_table(rows=5, cols=2)
        self._colorir_grid_tabela(info_table)
        self._set_table_full_width(info_table)

        info_table.cell(0, 0).text = "Cliente"
        info_table.cell(0, 1).text = cliente_nome
        info_table.cell(1, 0).text = "Perfil de Investidor"
        info_table.cell(1, 1).text = perfil
        info_table.cell(2, 0).text = "Data de Análise"
        info_table.cell(2, 1).text = datetime.now().strftime("%d/%m/%Y")
        info_table.cell(3, 0).text = "Patrimônio Total"
        info_table.cell(3, 1).text = f"R$ {patrimonio_total:,.2f}".replace(',', '.')
        info_table.cell(4, 0).text = "Número de Posições"
        info_table.cell(4, 1).text = str(len(posicoes))

        # Destacar a coluna de rótulos (col 0) com fundo cinza claro
        for row in info_table.rows:
            self._set_cell_background(row.cells[0], 'D6E4F0')
            for run in row.cells[0].paragraphs[0].runs:
                run.font.bold = True
        self._set_table_cell_margins(info_table)
        self._set_table_font_size(info_table, 9.5)
        self._remove_cell_spacing(info_table)
        self._colorir_negativos(info_table)

        # ===== RESUMO DE ALOCAÇÃO =====
        doc.add_paragraph()
        self._add_section_title(doc, "2. ALOCAÇÃO DA CARTEIRA")

        if patrimonio_total > 0:
            alocacao_table = doc.add_table(rows=len(alocacoes) + 1, cols=2)
            self._colorir_grid_tabela(alocacao_table)
            self._set_table_full_width(alocacao_table)

            alocacao_table.cell(0, 0).text = "Classe de Ativo"
            alocacao_table.cell(0, 1).text = "Alocação (%)"

            self._colorir_header_tabela(alocacao_table)
            self._set_alternating_rows(alocacao_table)

            for idx, (classe, pct) in enumerate(sorted(alocacoes.items(), key=lambda x: x[1], reverse=True), 1):
                alocacao_table.cell(idx, 0).text = classe
                alocacao_table.cell(idx, 1).text = f"{pct*100:.2f}%"

            self._set_table_cell_margins(alocacao_table)
            self._set_table_font_size(alocacao_table, 9.5)
            self._align_column(alocacao_table, 1)
            self._remove_cell_spacing(alocacao_table)
            self._colorir_negativos(alocacao_table)

        # ===== ANÁLISE DE SUITABILITY =====
        doc.add_paragraph()
        self._add_section_title(doc, "3. ANÁLISE DE ADERÊNCIA AO PERFIL")

        dentro_count = sum(1 for x in analise_suit if x['dentro_perfil'])
        fora_count = len(analise_suit) - dentro_count

        resumo = doc.add_paragraph()
        resumo.add_run(f"Dentro do Perfil: ").bold = True
        resumo.add_run(f"{dentro_count} classes | ")
        resumo.add_run(f"Fora do Perfil: ").bold = True
        resumo.add_run(f"{fora_count} classes")

        suit_table = doc.add_table(rows=len(analise_suit) + 1, cols=5)
        self._colorir_grid_tabela(suit_table)
        self._set_table_full_width(suit_table)

        headers = ['Classe de Ativo', 'Alocação Atual', 'Mín.', 'Máx.', 'Status']
        for col, header in enumerate(headers):
            suit_table.cell(0, col).text = header

        self._colorir_header_tabela(suit_table)
        self._set_alternating_rows(suit_table)

        for idx, item in enumerate(analise_suit, 1):
            suit_table.cell(idx, 0).text = item['classe']
            suit_table.cell(idx, 1).text = f"{item['alocacao_atual']*100:.2f}%"
            suit_table.cell(idx, 2).text = f"{item['limite_min']*100:.0f}%"
            suit_table.cell(idx, 3).text = f"{item['limite_max']*100:.0f}%"

            status_cell = suit_table.cell(idx, 4)
            if item['dentro_perfil']:
                status_cell.text = "✓ Dentro"
                self._set_cell_background(status_cell, 'C6EFCE')
            else:
                status_cell.text = "✗ Fora"
                self._set_cell_background(status_cell, 'FFC7CE')

        self._set_table_cell_margins(suit_table)
        self._set_table_font_size(suit_table, 9.5)
        self._align_column(suit_table, 1)
        self._align_column(suit_table, 2)
        self._align_column(suit_table, 3)
        self._align_column(suit_table, 4, WD_ALIGN_PARAGRAPH.CENTER)
        self._remove_cell_spacing(suit_table)
        self._colorir_negativos(suit_table)

        # ===== DETALHAMENTO E LIQUIDEZ =====
        doc.add_paragraph()
        self._add_section_title(doc, "4. DETALHAMENTO E LIQUIDEZ")

        nota = doc.add_paragraph()
        nota_run = nota.add_run(
            "Fundos: cotização conforme Guia de Fundos XP (mai/2026). "
            "Renda Fixa: vencimento do título. "
            "Ações e ETFs B3 (RV): D+2. "
            "Fundos Listados: D+2 (balcão). "
            "Alternativos (FIPs): sem liquidez determinada."
        )
        nota_run.font.italic = True
        nota_run.font.color.rgb = RGBColor(89, 89, 89)
        nota_run.font.size = Pt(8.5)

        # 4.1 DETALHES DAS POSIÇÕES
        subtitle_41 = doc.add_paragraph("4.1 DETALHES DAS POSIÇÕES")
        for run in subtitle_41.runs:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(27, 58, 92)

        pos_table = doc.add_table(rows=len(posicoes) + 1, cols=5)
        self._colorir_grid_tabela(pos_table)
        self._set_table_full_width(pos_table)

        headers = ['Ativo', 'Classe', 'Valor (R$)', 'Alocação', 'Liquidez/Prazo']
        for col, header in enumerate(headers):
            pos_table.cell(0, col).text = header

        self._colorir_header_tabela(pos_table)
        self._set_alternating_rows(pos_table)

        for idx, pos in enumerate(posicoes, 1):
            security_name = pos.get('security', {}).get('name', 'N/A')
            classe = pos.get('classe_ativo', 'N/A')
            valor = pos.get('valor_mercado', 0)
            alocacao = (valor / patrimonio_total * 100) if patrimonio_total > 0 else 0
            security = pos.get('security', {})
            liquidez = self.classificar_liquidez(security)

            pos_table.cell(idx, 0).text = security_name[:40]
            pos_table.cell(idx, 1).text = classe
            pos_table.cell(idx, 2).text = f"{valor:,.0f}".replace(',', '.')
            pos_table.cell(idx, 3).text = f"{alocacao:.2f}%"
            pos_table.cell(idx, 4).text = liquidez

        self._set_table_cell_margins(pos_table, top=40, left=80, bottom=40, right=80)
        self._set_table_font_size(pos_table, 7.5)
        self._align_column(pos_table, 2)
        self._align_column(pos_table, 3)
        self._remove_cell_spacing(pos_table)
        self._colorir_negativos(pos_table)

        # ===== ANÁLISE DE LIQUIDEZ =====
        doc.add_paragraph()
        self._add_section_title(doc, "4.2 ANÁLISE DE LIQUIDEZ DA CARTEIRA")

        # Classificar posições por liquidez
        liquidez_map = {}
        categorias = [
            "Disponível (D+0 / D+1)",
            "Curto Prazo (D+2 a D+30)",
            "Médio Prazo (D+31 a D+90)",
            "RF — vencimento determinado",
            "Sem liquidez (ilíquido)"
        ]

        for cat in categorias:
            liquidez_map[cat] = {'valor': 0, 'percentual': 0, 'posicoes': 0}

        for pos in posicoes:
            security = pos.get('security', {})
            valor = pos.get('valor_mercado', 0)
            categoria = self.classificar_liquidez(security)
            liquidez_map[categoria]['valor'] += valor
            liquidez_map[categoria]['posicoes'] += 1

        if patrimonio_total > 0:
            for categoria in liquidez_map:
                liquidez_map[categoria]['percentual'] = liquidez_map[categoria]['valor'] / patrimonio_total

        liquidez_table = doc.add_table(rows=len(liquidez_map) + 1, cols=4)
        self._colorir_grid_tabela(liquidez_table)
        self._set_table_full_width(liquidez_table)

        headers_liq = ['Categoria de Liquidez', 'Valor (R$)', 'Alocação %', 'Ativos']
        for col, header in enumerate(headers_liq):
            liquidez_table.cell(0, col).text = header

        self._colorir_header_tabela(liquidez_table)
        self._set_alternating_rows(liquidez_table)

        for idx, (categoria, dados) in enumerate(liquidez_map.items(), 1):
            valor = dados['valor']
            percentual = dados['percentual'] * 100
            num_posicoes = dados['posicoes']

            liquidez_table.cell(idx, 0).text = categoria
            liquidez_table.cell(idx, 1).text = f"R$ {valor:,.0f}".replace(',', '.')
            liquidez_table.cell(idx, 2).text = f"{percentual:.2f}%"
            liquidez_table.cell(idx, 3).text = str(num_posicoes)

        self._set_table_cell_margins(liquidez_table)
        self._set_table_font_size(liquidez_table, 9.5)
        self._align_column(liquidez_table, 1)
        self._align_column(liquidez_table, 2)
        self._align_column(liquidez_table, 3, WD_ALIGN_PARAGRAPH.CENTER)
        self._remove_cell_spacing(liquidez_table)
        self._colorir_negativos(liquidez_table)

        # ===== CONCLUSÕES =====
        doc.add_paragraph()
        self._add_section_title(doc, "5. CONCLUSÃO")

        if fora_count == 0:
            conclusao_texto = f"A carteira de {cliente_nome} apresenta total aderência ao perfil de investidor \'{perfil}\', com todas as classes de ativos dentro dos limites estabelecidos."
        else:
            conclusao_texto = f"A carteira de {cliente_nome} apresenta {fora_count} classe(s) de ativo(s) fora dos limites do perfil \'{perfil}\'. Recomenda-se rebalanceamento."

        conclusao_para = doc.add_paragraph(conclusao_texto)
        conclusao_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # ===== FOOTER =====
        self._add_footer(doc)

        # Salvar documento
        doc.save(output_path)
        logger.info(f"Laudo gerado com sucesso: {output_path}")

        return output_path

    @staticmethod
    def _set_document_font(doc, font_name: str):
        """Define a fonte padrão do documento inteiro"""
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE
        styles_el = doc.styles.element
        # Setar nos docDefaults
        rPrDef = styles_el.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPrDefault')
        if rPrDef is None:
            doc_defaults = styles_el.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}docDefaults')
            if doc_defaults is None:
                doc_defaults = _OE('w:docDefaults')
                styles_el.insert(0, doc_defaults)
            rPrDef = _OE('w:rPrDefault')
            doc_defaults.append(rPrDef)
        rPr = rPrDef.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
        if rPr is None:
            rPr = _OE('w:rPr')
            rPrDef.append(rPr)
        rFonts = rPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
        if rFonts is None:
            rFonts = _OE('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(_qn('w:ascii'), font_name)
        rFonts.set(_qn('w:hAnsi'), font_name)
        rFonts.set(_qn('w:cs'), font_name)
        # Setar também no estilo Normal
        try:
            doc.styles['Normal'].font.name = font_name
        except Exception:
            pass

    @staticmethod
    def _add_header(doc):
        """Adiciona cabeçalho ao documento"""
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "Fatorial Capital - Análise de Carteira"
        header_para.runs[0].font.size = Pt(10)
        header_para.runs[0].font.italic = True

    @staticmethod
    def _add_footer(doc):
        """Adiciona rodapé ao documento"""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')} - Fatorial Capital"
        footer_para.runs[0].font.size = Pt(8)
        footer_para.runs[0].font.italic = True

    @staticmethod
    def _add_section_title(doc, title: str):
        """Adiciona título de seção formatado"""
        p = doc.add_paragraph(title)
        p_format = p.paragraph_format
        p_format.space_before = Pt(6)
        p_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(27, 58, 92)

    @staticmethod
    def _set_cell_background(cell, fill):
        """Define cor de fundo de uma célula da tabela"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill)
        cell._element.get_or_add_tcPr().append(shading_elm)

    @staticmethod
    def _set_table_font_size(table, size_pt: int):
        """Define tamanho da fonte para toda a tabela"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(size_pt)

    @staticmethod
    def _colorir_negativos(table):
        """Colore valores negativos em vermelho em toda a tabela"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text.startswith('-'):
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)

    @staticmethod
    def _colorir_header_tabela(table):
        """Colore o header com fundo azul escuro, texto branco e centralizado"""
        for cell in table.rows[0].cells:
            GorilaLaudo._set_cell_background(cell, '1B3A5C')
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

    @staticmethod
    def _colorir_grid_tabela(table):
        """Bordas externas (1pt azul escuro) e internas (0.75pt azul escuro)"""
        DARK_BLUE = "1B3A5C"
        tbl = table._element
        tblPr = tbl.tblPr
        # Remover estilo pré-definido que sobrescreve formatação manual
        from docx.oxml.ns import qn as _qn
        tbl_style = tblPr.find(_qn('w:tblStyle')) if tblPr is not None else None
        if tbl_style is not None:
            tblPr.remove(tbl_style)
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        border_config = {
            'top':     (DARK_BLUE, '8'),
            'left':    (DARK_BLUE, '8'),
            'bottom':  (DARK_BLUE, '8'),
            'right':   (DARK_BLUE, '8'),
            'insideH': (DARK_BLUE, '6'),
            'insideV': (DARK_BLUE, '6'),
        }
        for border_name, (color, sz) in border_config.items():
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), sz)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            tblBorders.append(border)
        tblPr.append(tblBorders)

    @staticmethod
    def _set_table_full_width(table):
        """Faz a tabela ocupar 100% da largura da página"""
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        for old in tblPr.findall(qn('w:tblW')):
            tblPr.remove(old)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        tblPr.append(tblW)

    @staticmethod
    def _set_table_cell_margins(table, top=60, left=120, bottom=60, right=120):
        """Define margens internas (padding) das células"""
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblCellMar = OxmlElement('w:tblCellMar')
        for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:w'), str(val))
            el.set(qn('w:type'), 'dxa')
            tblCellMar.append(el)
        tblPr.append(tblCellMar)

    @staticmethod
    def _set_alternating_rows(table, color='DEEAF1'):
        """Zebra striping: linhas pares recebem fundo azul claro"""
        for idx, row in enumerate(table.rows[1:], 1):
            if idx % 2 == 0:
                for cell in row.cells:
                    GorilaLaudo._set_cell_background(cell, color)

    @staticmethod
    def _align_column(table, col_idx, alignment=WD_ALIGN_PARAGRAPH.RIGHT, skip_header=True):
        """Alinha o conteúdo de uma coluna específica"""
        start = 1 if skip_header else 0
        for row in table.rows[start:]:
            cell = row.cells[col_idx]
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignment

    @staticmethod
    def _remove_cell_spacing(table):
        """Remove espaçamento antes/depois dos parágrafos em todas as células"""
        from docx.shared import Pt as _Pt
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fmt = paragraph.paragraph_format
                    fmt.space_before = _Pt(0)
                    fmt.space_after = _Pt(0)


# ============================================================================
# FUNCAO PRINCIPAL
# ============================================================================

def main():
    """Funcao principal"""
    print("\n" + "="*80)
    print("GERADOR DE LAUDO DE CARTEIRA - API GORILA")
    print("="*80 + "\n")

    cliente_nome = 'RAPHAEL SANTOS DE ALMEIDA REZENDE DE MATTOS'  # input("Nome do cliente: ").strip()
    print("\nPerfis disponíveis: Conservador, Moderado, Agressivo")
    perfil = 'Moderado'  # input("Perfil de investidor: ").strip()

    if perfil not in ['Conservador', 'Moderado', 'Agressivo']:
        print("Erro: Perfil inválido!")
        sys.exit(1)

    try:
        system = GorilaLaudo(".env")

        portfolios = system.buscar_portfolios()

        if not portfolios:
            print("Erro: Nenhum port